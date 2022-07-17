import os
import tkinter as tk
import datetime
import docx

job_window=tk.Tk()
job_window.title('Job Application Tracker')
job_window.config(bg='grey')

#To add,
#  Show Example Results or Show local results (keep local results on home directory so that its private)
#   Add a counter to count the amount of psoitions applied for
#   Add shortcut buttons for certain 
#   Add a text box to assist in cover letter creation
#   Add buttons to open Linked in and Indeed websites

path2data=os.path.dirname(os.path.dirname(os.path.dirname(__file__)))+'\\Local_Data'
path2jobs=path2data+'\\jobsapplied.txt'
if os.path.exists(path2data)==True:
    pass
else:
    print("No Local Data folder detected, If you're doing this from a clone of the GitHub, I am keeping my log files private")
    print("If you want to try them out, create a Local_Data folder in the same directory as Personal_Projects_Public directory")

emptydoc_path=os.path.dirname(__file__)+'\\EmptyDoc.docx'
rv=0
title_lab=tk.Label(text='Job Application Tracker',bg='grey',fg='Black').grid(row=rv,columnspan=8)
rv+=1
web_lab=tk.Label(text='Website :',bg='grey',fg='cyan').grid(row=rv,column=0)
web_ent=tk.Entry()
web_ent.grid(row=rv,column=1)
comp_lab=tk.Label(text='Company :',bg='grey',fg='cyan').grid(row=rv,column=2)
comp_ent=tk.Entry()
comp_ent.grid(row=rv,column=3)
pos_lab=tk.Label(text='Position :',bg='grey',fg='cyan').grid(row=rv,column=4)
pos_ent=tk.Entry()
pos_ent.grid(row=rv,column=5)
date_lab=tk.Label(text='Date :',bg='grey',fg='cyan').grid(row=rv,column=6)
date_ent=tk.Entry()
date_ent.grid(row=rv,column=7)

todaydate=datetime.date.today()
todaydate=str(todaydate)
date_ent.delete('0',tk.END)
date_ent.insert('0',todaydate)

rv+=1

def add_job():
    website=web_ent.get()
    company=comp_ent.get()
    position=pos_ent.get()
    date=str(date_ent.get())
    num_apps=num_app_ent.get()

    with open(path2jobs,"a") as textfile:
        textfile.write('{},{},{},{}\n'.format(website,company,position,date))
    
    num_apps=int(num_apps)
    num_apps+=1
    num_apps=str(num_apps)
    num_app_ent.delete('0',tk.END)
    num_app_ent.insert('0',num_apps)

def read_jobs():
    with open(path2jobs) as joblist:
        for line in joblist:
            linebits=line.split(',')
            linebits

def history():
    pass

def lnk():
    web_ent.delete('0',tk.END)
    web_ent.insert('0','LinkedIn')

def ind():
    web_ent.delete('0',tk.END)
    web_ent.insert('0','Indeed')

def comp_site():
    web_ent.delete('0',tk.END)
    web_ent.insert('0','CompanyWebsite')

def mech():
    pos_ent.delete('0',tk.END)
    pos_ent.insert('0','Mechanical Engineer')


if os.path.exists(path2jobs)==True:
    with open(path2jobs,'r') as file:
        linect=0
        for line in file:
            linect+=1
    # num_app_ent.delete('0',tk.END)
    # num_app_ent.insert('0',str(linect))
else:
    with open(path2jobs,'w') as file:
        pass

rv=6
add_app_button=tk.Button(text='Job Applied!',command=add_job,bg='black',fg='cyan',width=36).grid(row=rv,column=7)
#Shortcuts:
sc_lab=tk.Label(text='Shortcuts:',bg='grey').grid(column=0,row=rv)
linkedin_sc=tk.Button(text='LinkedIn',command=lnk,fg='navy',font='Helvetica').grid(column=1,row=rv)
indeed_sc=tk.Button(text='Indeed',command=lnk,fg='blue',font='Times').grid(column=2,row=rv)
comp_sc=tk.Button(text='Company Website',command=comp_site,fg='dark red').grid(column=3,row=rv)
mech_sc=tk.Button(text='Mechanical Engineer',command=mech).grid(column=4,row=rv)

cover_textbox=tk.Text()
cover_textbox.grid(columnspan=8)
cover_textbox.delete('1.0',tk.END)
cover_textbox.insert('1.0','<Enter message Here>')

def gen_cover():
    # website=web_ent.get()
    company=comp_ent.get()
    # position=pos_ent.get()
    date=str(date_ent.get())
    covertext=cover_textbox.get('1.0',tk.END)
    docname='BDyer_CoverLetter_{}_{}.docx'.format(company,date)
    docpath=path2data+'\\'+docname
    if os.path.exists(docpath):
        pass
    else:
        with open(docpath,"w") as file:
            pass
    mydoc=docx.Document(docpath)
    mydoc.add_paragraph(date)
    mydoc.add_paragraph("")
    mydoc.add_paragraph("Benjamin Dyer")
    mydoc.add_paragraph("2527 Owens Ave, Unit 204")
    mydoc.add_paragraph("Fort Collins, CO, 80528")
    mydoc.add_paragraph("720-320-1601")
    mydoc.add_paragraph("bdyer5280@gmail.com")
    mydoc.add_paragraph("")
    mydoc.add_paragraph("   Hello,")
    mydoc.add_paragraph("")
    mydoc.add_paragraph(covertext)
    mydoc.add_paragraph("")
    mydoc.add_paragraph("Thank you for your consideration,")
    mydoc.add_paragraph("")    
    mydoc.add_paragraph("Benjamin Dyer")
    
    mydoc.save(docpath)
    

generate_cover_butt=tk.Button(text='Generate Cover Letter',command=gen_cover).grid(columnspan=8)
num_app_lab=tk.Label(text='Number of positions applied to:',bg='grey',fg='black').grid()
num_app_ent=tk.Entry()
num_app_ent.grid()
jb_ct=0
with open(path2jobs,"r") as file:
    for line in file:
        if ',' in line:
            jb_ct+=1

num_app_ent.delete('0',tk.END)
num_app_ent.insert('0',jb_ct)

job_window.mainloop()
